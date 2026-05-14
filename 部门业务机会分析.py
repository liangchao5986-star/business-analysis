""" 销售漏斗智能分析工具 v1.0 """
import os, re, tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
import numpy as np
from datetime import datetime, timedelta
from collections import defaultdict
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ================================================================
# 常量
# ================================================================
STAGE_ORDER = ['销售线索', '公司立项', '客户立项', '技术认可', '商务认可', '签约准备', '已赢单', '已输单']
STAGE_WEIGHT = {
    '销售线索': 0.05, '公司立项': 0.10, '客户立项': 0.30, '技术认可': 0.50,
    '商务认可': 0.75, '签约准备': 1.00, '已赢单': 1.00, '已输单': 0.00
}
STAGE_WEIGHT_ORDERED = {s: STAGE_WEIGHT.get(s, 0.1) for s in STAGE_ORDER}
INVALID_REASON_MAP = {
    'ancient': '🗑️ 远古僵尸单：创建超过3年仍未签约',
    'not_2026': '🗑️ 非本年度项目：预计签约日期不在2026年',
    'stagnant_prep': '🗑️ 签约准备停滞：阶段修改日期距今>180天',
    'stagnant_tech': '🗑️ 技术认可停滞：阶段修改日期距今>360天',
    'stagnant_comm': '🗑️ 商务认可停滞：阶段修改日期距今>360天',
    'zero_amount': '🗑️ 金额为零：无法带来收入贡献',
    'zero_prob': '🗑️ 阶段概率不匹配：后期阶段可能性过低',
    'no_update': '🗑️ 超180天未更新：疑似已放弃跟进',
    'pure_borrow': '🗑️ 纯借机项目：无直接销售收入',
    'tiny_amount': '⚠️ 微型金额(<5000元)',
}
REQUIRED_COLS = ['阶段', '创建日期', '预计签约金额', '业务机会所有人', '销售类型', '更新时间', '业务机会名', '项目类型', '可能性', '预计签约日期', '阶段修改日期']
ALL_STANDARD_NAMES = [
    '阶段', '预计签约金额', '业务机会所有人', '业务机会名', '创建日期', '销售类型', '更新时间', '项目类型', '可能性', '行业', '客户名', '销售部门', '产品描述', '预计签约日期', '未签报预计回款金额', '业务机会编号', '结果', '主要市场活动来源', '是否Name客户', '是否有定制开发', '已转签报单', '签约客户联系人', '最终客户', '最终客户联系人', '丢单原因', '竞争对手', '丢单原因备注', '价格手册', '客户编号', '预计供货日期', '是否预计签约当月验收', '阶段修改日期', '项目名称', '项目编号', '签报单流水号', '合同名称', '合同编号', '合计', '所有人编号', '所属公司', '原销售人员', '所有人部门', '创建人', '创建时间', '修改时间', '事业部',
]

# ================================================================
# 工具函数
# ================================================================
def safe_col(df, col_name, fallback=''):
    if col_name in df.columns:
        return df[col_name]
    return pd.Series([fallback] * len(df), index=df.index, name=col_name)

def format_amount(val):
    try:
        v = float(val)
        return f'{v / 10000:.2f}万' if abs(v) >= 10000 else f'{v:,.0f}元'
    except:
        return str(val) if val else '-'

def format_date(val):
    try:
        if pd.isna(val): return '-'
        if isinstance(val, (pd.Timestamp, datetime)): return val.strftime('%Y-%m-%d')
        return str(val)[:10]
    except:
        return '-'

def safe_timedelta(dt):
    if pd.isna(dt): return 0
    try:
        return (datetime.now() - pd.to_datetime(dt)).days
    except:
        return 0

def clean_col_name(name):
    s = str(name)
    s = re.sub(r'[\u200b\u200c\u200d\ufeff\u00a0\u200e\u200f]', '', s)
    s = s.strip()
    s = re.sub(r'\s+', ' ', s)
    return s

# ================================================================
# 智能Excel读取（纯 pandas，杜绝 ZIP closed）
# ================================================================
def smart_read_excel(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.csv':
        df = pd.read_csv(file_path)
        df.columns = [clean_col_name(c) for c in df.columns]
        return df, 0
    best_df = None
    best_header = 0
    best_score = 0
    for h in range(6):
        try:
            test_df = pd.read_excel(file_path, header=h)
            test_df.columns = [clean_col_name(c) for c in test_df.columns]
            score = sum(1 for c in test_df.columns if c in ALL_STANDARD_NAMES)
            if score > best_score:
                best_score = score
                best_header = h
                best_df = test_df
        except Exception:
            continue
    if best_df is not None and best_score >= 3:
        best_df.columns = [clean_col_name(c) for c in best_df.columns]
        return best_df, best_header
    try:
        df = pd.read_excel(file_path)
        df.columns = [clean_col_name(c) for c in df.columns]
        return df, 0
    except Exception:
        raise ValueError(f'无法读取Excel文件：{file_path}')

# ================================================================
# 列名映射
# ================================================================
def map_columns_enhanced(df):
    col_map = {}
    cols = list(df.columns)
    rules = {
        '业务机会名': ['业务机会名', '商机名称', '商机名', '项目名称', '项目名', '名称', '机会名称'],
        '创建日期': ['创建日期', '创建时间', '录入日期', '录入时间', '创建日', '建档日期'],
        '阶段': ['阶段', '商机阶段', '业务阶段', '当前阶段', '阶段名称', '所处阶段', '状态'],
        '预计签约金额': ['预计签约金额', '预计金额', '金额', '预计签约额', '签约金额', '预估金额'],
        '业务机会所有人': ['业务机会所有人', '负责人', '销售', '销售负责人', '所有人', '销售员', '客户经理', '销售人'],
        '业务机会编号': ['业务机会编号', '商机编号', '编号', 'OPP编号', '机会编号'],
        '行业': ['行业', '行业分类', '行业类别', '客户行业'],
        '销售类型': ['销售类型', '销售模式', '签约方式', '类型'],
        '更新时间': ['更新时间', '修改时间', '最后更新', '更新日期', '最后修改', '修改日期'],
        '预计签约日期': ['预计签约日期', '预计签约', '计划签约日期', '签约日期', '计划签约'],
        '未签报预计回款金额': ['未签报预计回款金额', '预计回款金额', '回款金额', '预计回款'],
        '项目类型': ['项目类型', '项目类别', '商机类型'],
        '客户名': ['客户名', '客户名称', '客户', '客户简称'],
        '销售部门': ['销售部门', '部门', '所属部门', '团队'],
        '产品描述': ['产品描述', '产品', '产品名称', '产品线'],
        '可能性': ['可能性', '赢率', '概率', '胜率', '成交概率', '预计赢率'],
        '阶段修改日期': ['阶段修改日期', '阶段变更日期', '阶段更新日期', '阶段推进日期'],
    }
    for standard, aliases in rules.items():
        if standard in cols: continue
        matched = False
        for alias in aliases:
            alias_str = str(alias) if not isinstance(alias, str) else alias
            for c in cols:
                if c.lower() == alias_str.lower():
                    col_map[c] = standard; matched = True; break
            if matched: break
        if not matched:
            for alias in aliases:
                alias_str = str(alias) if not isinstance(alias, str) else alias
                if len(alias_str) < 2: continue
                for c in cols:
                    if alias_str in c or c in alias_str:
                        col_map[c] = standard; matched = True; break
                if matched: break
        if not matched:
            for c in cols:
                if len(standard) >= 2 and standard in c:
                    col_map[c] = standard; break
    return col_map

# ================================================================
# 分析引擎
# ================================================================
class FunnelAnalyzer:
    def __init__(self, df):
        self.df = df.copy()
        self.df.columns = [clean_col_name(c) for c in self.df.columns]
        self._parse_dates()
        self._parse_amounts()
        self.invalid_records = []
        self.有效业务机会_filters = {}
        self.high_value = pd.DataFrame()
        self.all_valid = pd.DataFrame()
        self.summary = {}
        self.missing_cols = []
        self.read_diagnosis = {}
        self.excluded_by_year = 0

    def _parse_dates(self):
        for c in self.df.columns:
            for kw in ['日期', '时间', 'date', 'time']:
                if kw in c.lower():
                    self.df[c] = pd.to_datetime(self.df[c], errors='coerce')
                    break

    def _parse_amounts(self):
        for c in self.df.columns:
            for kw in ['金额', '回款', 'amount', '合计']:
                if kw in c.lower():
                    self.df[c] = pd.to_numeric(self.df[c], errors='coerce').fillna(0)
                    break

    def _ensure_col(self, col_name):
        if col_name not in self.df.columns:
            if col_name not in self.missing_cols:
                self.missing_cols.append(col_name)
            self.df[col_name] = ''

    def run_analysis(self):
        self.read_diagnosis = {
            'total_columns': len(self.df.columns),
            'column_names': list(self.df.columns),
        }
        for col in REQUIRED_COLS:
            self._ensure_col(col)
        self._step0_filter_year()
        self._step1()
        self._step2()
        self._step3()
        self._step4()
        self._step5()
        return self.summary

    def _step0_filter_year(self):
        """第0步：筛选预计签约日期为2026年的项目"""
        col = '预计签约日期'
        total_before = len(self.df)
        if col in self.df.columns:
            mask = self.df[col].notna()
            year_mask = pd.Series([False] * len(self.df), index=self.df.index)
            for idx in self.df.index:
                val = self.df.at[idx, col]
                if pd.notna(val):
                    try:
                        dt = pd.to_datetime(val)
                        if dt.year == 2026:
                            year_mask.at[idx] = True
                    except: pass
            excluded = mask & ~year_mask
            self.excluded_by_year = int(excluded.sum())
            self.df = self.df[year_mask | ~mask].copy().reset_index(drop=True)
        else:
            self.excluded_by_year = 0
        after = len(self.df)
        self.read_diagnosis['year_filter'] = {
            'before': total_before, 'excluded': self.excluded_by_year, 'after': after,
        }

    def _step1(self):
        self.invalid_records = []
        self.df['_invalid_reasons'] = [[] for _ in range(len(self.df))]
        for i, row in self.df.iterrows():
            reasons = []
            create_dt = row.get('创建日期', pd.NaT)
            update_dt = row.get('更新时间', pd.NaT)
            stage_mod_dt = row.get('阶段修改日期', pd.NaT)
            stage = str(row.get('阶段', '') or '').strip()
            amt = float(row.get('预计签约金额', 0) or 0)
            proj_type = str(row.get('项目类型', '') or '')
            prob_str = str(row.get('可能性', '') or '')
            age = safe_timedelta(create_dt)
            if age > 365 * 3: reasons.append('ancient')
            stag_age = safe_timedelta(stage_mod_dt)
            if stag_age == 0: stag_age = safe_timedelta(update_dt)
            if stag_age == 0: stag_age = age
            if stage == '签约准备' and stag_age > 180: reasons.append('stagnant_prep')
            elif stage == '技术认可' and stag_age > 360: reasons.append('stagnant_tech')
            elif stage == '商务认可' and stag_age > 360: reasons.append('stagnant_comm')
            if amt == 0: reasons.append('zero_amount')
            if stage in ('商务认可', '签约准备'):
                try:
                    p = float(prob_str.replace('%', '')) if '%' in prob_str else float(prob_str)
                    if stage == '签约准备' and p < 90: reasons.append('zero_prob')
                except: pass
            if safe_timedelta(update_dt) > 180: reasons.append('no_update')
            if '纯借机' in proj_type: reasons.append('pure_borrow')
            if 0 < amt < 5000: reasons.append('tiny_amount')
            self.df.at[i, '_invalid_reasons'] = reasons
            for r in reasons:
                self.invalid_records.append({
                    'index': i, 'name': str(row.get('业务机会名', f'行{i}') or '')[:50],
                    'stage': stage, 'owner': str(row.get('业务机会所有人', '') or ''),
                    'amount': amt, 'create_date': create_dt, 'stage_mod_date': stage_mod_dt,
                    'reason_code': r, 'reason_text': INVALID_REASON_MAP.get(r, r),
                })
        self.df['_is_invalid'] = self.df['_invalid_reasons'].apply(lambda x: len(x) > 0)

    def _step2(self):
        self.all_valid = self.df[~self.df['_is_invalid']].copy()

    def _step3(self):
        df = self.all_valid.copy()
        steps = {'原始有效机会': df.copy()}
        s_s = safe_col(df, '阶段', '')
        s1 = df[~s_s.isin(['销售线索', '公司立项'])].copy()
        steps['甄别S1｜排除销售线索与公司立项'] = s1
        s1s = safe_col(s1, '阶段', '')
        s2 = s1[s1s.isin(['客户立项', '技术认可', '商务认可', '签约准备'])].copy()
        steps['甄别S2｜保留客户立项及以后'] = s2
        st = safe_col(s2, '销售类型', '')
        m3 = st.isin(['直签用户', '渠道销售', ''])
        s3 = s2[m3].copy()
        s3['_is_direct'] = (st[m3] == '直签用户')
        steps['甄别S3｜排除不可控渠道类型'] = s3
        uc = safe_col(s3, '更新时间', pd.NaT)
        if uc.notna().any():
            cutoff = datetime.now() - timedelta(days=120)
            try:
                m4 = uc.notna() & (pd.to_datetime(uc, errors='coerce') >= cutoff)
            except:
                m4 = pd.Series([True] * len(s3), index=s3.index)
            s4 = s3[m4].copy()
        else:
            s4 = s3.copy()
        steps['甄别S4｜120天内有活跃更新'] = s4
        ac = pd.to_numeric(safe_col(s4, '预计签约金额', 0), errors='coerce').fillna(0)
        s5 = s4[ac > 5000].copy()
        steps['甄别S5｜有效金额>5000元'] = s5
        self.有效业务机会_filters = steps

    def _step4(self):
        key = '甄别S5｜有效金额>5000元'
        if key not in self.有效业务机会_filters:
            self.high_value = pd.DataFrame(); return
        df = self.有效业务机会_filters[key].copy()
        if df.empty: self.high_value = pd.DataFrame(); return
        ss = safe_col(df, '阶段', '')
        an = pd.to_numeric(safe_col(df, '预计签约金额', 0), errors='coerce').fillna(0)
        df['_weight'] = ss.map(STAGE_WEIGHT_ORDERED).fillna(0.1)
        df['_amt'] = an
        df['_weighted_amt'] = an * df['_weight']
        df['_score_stage'] = df['_weight'] * 50
        df['_score_amount'] = 0
        df.loc[an >= 500000, '_score_amount'] = 30
        df.loc[(an >= 100000) & (an < 500000), '_score_amount'] = 20
        df.loc[(an >= 30000) & (an < 100000), '_score_amount'] = 10
        df.loc[(an > 5000) & (an < 30000), '_score_amount'] = 5
        df['_score_direct'] = 0
        if '_is_direct' in df.columns:
            df.loc[df['_is_direct'], '_score_direct'] = 20
        df['_total_score'] = df['_score_stage'] + df['_score_amount'] + df['_score_direct']
        self.high_value = df.sort_values( ['_total_score', '_weighted_amt'], ascending=[False, False] ).reset_index(drop=True)

    def _step5(self):
        total = len(self.df)
        ic = int(self.df['_is_invalid'].sum()) if '_is_invalid' in self.df.columns else 0
        vc = total - ic
        aa = pd.to_numeric(safe_col(self.df, '预计签约金额', 0), errors='coerce').fillna(0)
        ta = float(aa.sum())
        ia = float(aa[self.df['_is_invalid']].sum()) if ic > 0 else 0
        va = ta - ia
        hvc = len(self.high_value)
        hva = float(self.high_value['_amt'].sum()) if not self.high_value.empty else 0
        hvw = float(self.high_value['_weighted_amt'].sum()) if not self.high_value.empty else 0
        sd = {}
        ss = safe_col(self.df, '阶段', '')
        for s in STAGE_ORDER:
            m = ss == s; c = int(m.sum())
            if c > 0:
                sd[s] = {'count': c, 'amount': float(aa[m].sum())}
        od = {}
        if '业务机会所有人' in self.all_valid.columns and not self.all_valid.empty:
            for n, g in self.all_valid.groupby('业务机会所有人'):
                a = pd.to_numeric(safe_col(g, '预计签约金额', 0), errors='coerce').fillna(0)
                w = safe_col(g, '阶段', '').map(STAGE_WEIGHT_ORDERED).fillna(0.1)
                od[str(n)] = {
                    'count': len(g), 'amount': float(a.sum()), 'weighted': float((a * w).sum()),
                }
        self.summary = {
            'total_count': total, 'invalid_count': ic, 'valid_count': vc,
            'total_amount': ta, 'invalid_amount': ia, 'valid_amount': va,
            'high_value_count': hvc, 'high_value_amount': hva, 'high_value_weighted': hvw,
            'stage_distribution': sd, 'owner_distribution': od,
            'missing_cols': self.missing_cols, 'read_diagnosis': self.read_diagnosis,
            'excluded_by_year': self.excluded_by_year,
        }

# ================================================================
# 图表生成
# ================================================================
def generate_charts(analyzer, save_dir):
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    plt.rcParams['font.sans-serif'] = [
        'SimHei', 'Microsoft YaHei', 'WenQuanYi Micro Hei', 'Arial Unicode MS', 'Noto Sans CJK SC', 'DejaVu Sans'
    ]
    plt.rcParams['axes.unicode_minus'] = False
    os.makedirs(save_dir, exist_ok=True)
    cf = []
    ss = analyzer.summary.get('stage_distribution', {})
    if ss:
        fig, ax = plt.subplots(figsize=(10, 5))
        stages = [s for s in STAGE_ORDER if s in ss]
        amounts = [ss[s]['amount'] / 10000 for s in stages]
        counts = [ss[s]['count'] for s in stages]
        colors_list = ['#95a5a6', '#bdc3c7', '#3498db', '#2ecc71', '#f39c12', '#e74c3c']
        bar_colors = []
        for i in range(len(stages)):
            idx = min(i, len(colors_list) - 1)
            bar_colors.append(colors_list[idx])
        bars = ax.barh(range(len(stages)), amounts, color=bar_colors, edgecolor='white', height=0.6)
        for b, c, a in zip(bars, counts, amounts):
            ax.text(b.get_width() + max(amounts) * 0.01, b.get_y() + b.get_height() / 2, f'{a:.1f}万({c}单)', va='center', fontsize=9)
        ax.set_yticks(range(len(stages)))
        ax.set_yticklabels(stages, fontsize=10)
        ax.set_xlabel('金额(万元)')
        ax.set_title('销售漏斗各阶段金额分布', fontsize=13, fontweight='bold')
        ax.invert_yaxis()
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        plt.tight_layout()
        p = os.path.join(save_dir, 'chart_stage_distribution.png')
        plt.savefig(p, dpi=150, bbox_inches='tight')
        plt.close()
        cf.append(p)

    if analyzer.invalid_records:
        fig, ax = plt.subplots(figsize=(9, 6))
        rc = {}
        for r in analyzer.invalid_records:
            rc[r['reason_code']] = rc.get(r['reason_code'], 0) + 1
        labels = []
        for k in rc.keys():
            val = str(INVALID_REASON_MAP.get(k, k))
            parts = val.split('：', 1)
            labels.append(parts[0] if parts else val)
        ax.pie(list(rc.values()), labels=labels, autopct='%1.1f%%', explode=[0.05] * len(rc), startangle=90, textprops={'fontsize': 8}, pctdistance=0.85)
        ax.set_title(f'无效机会原因分类(共{len(analyzer.invalid_records)}条)', fontsize=13, fontweight='bold')
        plt.tight_layout()
        p = os.path.join(save_dir, 'chart_invalid_reasons.png')
        plt.savefig(p, dpi=150, bbox_inches='tight')
        plt.close()
        cf.append(p)

    od = analyzer.summary.get('owner_distribution', {})
    if od:
        fig, ax = plt.subplots(figsize=(12, 6))
        so = sorted(od.items(), key=lambda x: -x[1]['amount'])[:15]
        names = [o[0] for o in so]
        nom = [o[1]['amount'] / 10000 for o in so]
        wgt = [o[1]['weighted'] / 10000 for o in so]
        x = range(len(names))
        bw = 0.35
        ax.bar([i - bw / 2 for i in x], nom, bw, label='名义金额', color='#3498db', alpha=0.8)
        ax.bar([i + bw / 2 for i in x], wgt, bw, label='加权金额', color='#e74c3c', alpha=0.8)
        ax.set_xticks(x)
        ax.set_xticklabels(names, rotation=45, ha='right', fontsize=16)
        ax.set_ylabel('金额(万元)', fontsize=14)
        ax.set_title('销售人员有效漏斗对比', fontsize=16, fontweight='bold')
        ax.legend(fontsize=18)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        plt.tight_layout()
        p = os.path.join(save_dir, 'chart_owner_comparison.png')
        plt.savefig(p, dpi=150, bbox_inches='tight')
        plt.close()
        cf.append(p)

    if not analyzer.high_value.empty:
        top = analyzer.high_value.head(20).copy()
        fig, ax = plt.subplots(figsize=(12, 7))
        names = [str(n)[:30] for n in top['业务机会名'].tolist()]
        nom = top['_amt'].values / 10000
        wgt = top['_weighted_amt'].values / 10000
        sc = top['_total_score'].values
        y = range(len(names))
        ax.barh([i - 0.2 for i in y], nom, 0.4, label='名义金额', color='#3498db', alpha=0.8)
        ax.barh([i + 0.2 for i in y], wgt, 0.4, label='加权金额', color='#27ae60', alpha=0.8)
        for i, s in enumerate(sc):
            ax.text(max(nom[i], wgt[i]) + max(nom.max(), wgt.max()) * 0.02, i, f'{s:.0f}分', va='center', fontsize=14, color='#e74c3c', fontweight='bold')
        ax.set_yticks(y)
        ax.set_yticklabels(names, fontsize=14)
        ax.set_xlabel('金额(万元)', fontsize=14)
        ax.set_title('高价值机会Top20(含综合评分)', fontsize=16, fontweight='bold')
        ax.invert_yaxis()
        ax.legend(fontsize=18)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        plt.tight_layout()
        p = os.path.join(save_dir, 'chart_high_value_top20.png')
        plt.savefig(p, dpi=150, bbox_inches='tight')
        plt.close()
        cf.append(p)

    fig, ax = plt.subplots(figsize=(10, 5))
    fn = list(analyzer.有效业务机会_filters.keys())
    fc_vals = [len(d) for d in analyzer.有效业务机会_filters.values()]
    ax.fill_between(range(len(fn)), fc_vals, alpha=0.3, color='#3498db')
    ax.plot(range(len(fn)), fc_vals, 'o-', color='#2c3e50', linewidth=2, markersize=8)
    for i, (n, c) in enumerate(zip(fn, fc_vals)):
        ax.annotate(str(c), (i, c), textcoords="offset points", xytext=(0, 12), ha='center', fontsize=18, fontweight='bold')
    ax.set_xticks(range(len(fn)))
    ax.set_xticklabels( [n.split('｜')[-1] if '｜' in n else n for n in fn], rotation=25, ha='right', fontsize=14 )
    ax.set_ylabel('机会数量', fontsize=14)
    ax.set_title('有效业务机会甄别过程漏斗', fontsize=16, fontweight='bold')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    p = os.path.join(save_dir, 'chart_有效业务机会_funnel.png')
    plt.savefig(p, dpi=150, bbox_inches='tight')
    plt.close()
    cf.append(p)

    return cf

# ================================================================
# DOCX报告生成
# ================================================================
def set_cell_shading(cell, color):
    cell._tc.get_or_add_tcPr().append( parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>') )

def set_run_font(run, font_name='宋体', size=None, bold=None, color=None):
    run.font.name = font_name
    r = run._element
    rPr = r.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font_name)
    if size is not None: run.font.size = size
    if bold is not None: run.bold = bold
    if color is not None: run.font.color.rgb = color

def remove_empty_paragraph_before_picture(doc):
    """清除文档中紧跟在分页符或标题后面的空段落，彻底解决图表导致的空白页问题"""
    paras_to_remove = []
    for i, para in enumerate(doc.paragraphs):
        is_break_or_heading = False
        # 检查是否为分页符段落
        for run in para.runs:
            if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                is_break_or_heading = True
                break
        # 检查是否为标题 (通过样式判断)
        if not is_break_or_heading and para.style.name.startswith('Heading'):
            is_break_or_heading = True
        if is_break_or_heading and i + 1 < len(doc.paragraphs):
            next_para = doc.paragraphs[i + 1]
            # 如果下一个段落完全没有文字且没有任何Run，说明是插入图片留下的空壳
            if next_para.text == '' and not next_para.runs:
                paras_to_remove.append(next_para._element)
    for p_element in paras_to_remove:
        p_element.getparent().remove(p_element)

def add_styled_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, font_name='宋体', size=Pt(8), bold=True, color=RGBColor(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2C3E50')
    for ri, rd in enumerate(data):
        for ci, val in enumerate(rd):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            set_run_font(r, font_name='宋体', size=Pt(7.5))
            if ri % 2 == 0:
                set_cell_shading(cell, 'EBF5FB')
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table

def generate_report(analyzer, chart_files, output_path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2); section.bottom_margin = Cm(2)
    section.left_margin = Cm(2); section.right_margin = Cm(2)

    # 修改默认样式的字体为宋体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

    for _ in range(4): doc.add_paragraph('')
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('销售漏斗智能分析报告')
    set_run_font(r, font_name='宋体', size=Pt(28), bold=True, color=RGBColor(44, 62, 80))
    st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run('基于有效业务机会模型与漏斗健康度理论')
    set_run_font(r, font_name='宋体', size=Pt(14), color=RGBColor(127, 140, 141))
    doc.add_paragraph('')
    dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = dp.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
    set_run_font(r, font_name='宋体', size=Pt(11), color=RGBColor(149, 165, 166))

    diag = analyzer.read_diagnosis
    yf = diag.get('year_filter', {})
    if yf and yf.get('excluded', 0) > 0:
        doc.add_paragraph('')
        wp = doc.add_paragraph(); wp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = wp.add_run( f'📅 已筛选2026年签约项目：原始{yf["before"]}条，' f'排除{yf["excluded"]}条非2026年项目，保留{yf["after"]}条' )
        set_run_font(r, font_name='宋体', size=Pt(9), color=RGBColor(52, 152, 219))

    if analyzer.missing_cols:
        doc.add_paragraph('')
        wp2 = doc.add_paragraph(); wp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = wp2.add_run(f'⚠️ 未识别列：{", ".join(analyzer.missing_cols)}')
        set_run_font(r, font_name='宋体', size=Pt(9), color=RGBColor(231, 76, 60))

    doc.add_paragraph('')
    dp2 = doc.add_paragraph(); dp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = dp2.add_run( f'📊 当前分析：{diag.get("total_columns", 0)}列，' f'{analyzer.summary["total_count"]}条记录' )
    set_run_font(r, font_name='宋体', size=Pt(9), color=RGBColor(52, 152, 219))

    doc.add_page_break()
    h1 = doc.add_heading('目 录', level=1)
    for run in h1.runs: set_run_font(run, font_name='宋体')
    for item in ['一、数据总览与核心指标', '二、漏斗健康度诊断', '三、无效业务机会识别与剔除', '四、有效业务机会甄别与高价值机会清单', '五、人员漏斗透视', '六、管理行动建议']:
        p = doc.add_paragraph()
        r = p.add_run(item)
        set_run_font(r, font_name='宋体', size=Pt(11))

    doc.add_page_break()
    s = analyzer.summary

    # ---- 第一章 ----
    h1 = doc.add_heading('一、数据总览与核心指标', level=1)
    for run in h1.runs: set_run_font(run, font_name='宋体')
    h2 = doc.add_heading('1.1 核心数据指标', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    exc_yr = analyzer.summary.get('excluded_by_year', 0)
    kpi = [
        ['原始记录数', f'{yf.get("before", 0)}条', 'Excel全部记录'],
        ['非2026年排除', f'{exc_yr}条', '预计签约日期非2026年'],
        ['本次分析数', f'{s["total_count"]}条', '2026年签约项目'],
        ['总预计签约金额', format_amount(s['total_amount']), '2026年项目金额合计'],
        ['无效机会数', f'{s["invalid_count"]}个', f'占比{s["invalid_count"] / max(s["total_count"], 1) * 100:.1f}%'],
        ['无效机会金额', format_amount(s['invalid_amount']), '被剔除金额'],
        ['有效机会数', f'{s["valid_count"]}个', f'占比{s["valid_count"] / max(s["total_count"], 1) * 100:.1f}%'],
        ['有效机会金额', format_amount(s['valid_amount']), '有效漏斗金额'],
        ['高价值机会数', f'{s["high_value_count"]}个', '有效业务机会甄别真机会'],
        ['高价值加权金额', format_amount(s['high_value_weighted']), '阶段概率折算收入'],
    ]
    add_styled_table(doc, ['指标项', '数值', '说明'], kpi, [4, 4, 8])

    h2 = doc.add_heading('1.2 漏斗阶段分布', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    stage_stats = s.get('stage_distribution', {})
    if stage_stats:
        ta = sum(v['amount'] for v in stage_stats.values())
        sd_rows = []
        for sn in STAGE_ORDER:
            if sn in stage_stats:
                v = stage_stats[sn]
                w = STAGE_WEIGHT_ORDERED.get(sn, 0)
                wa = v['amount'] * w
                pct = v['amount'] / max(ta, 1) * 100
                note = ''
                if pct > 30 and sn in ('客户立项', '技术认可'): note = '⚠️ 大肚腩'
                if sn == '签约准备' and pct < 5 and ta > 0: note = '⚠️ 出口过窄'
                sd_rows.append([sn, str(v['count']), format_amount(v['amount']), f'{pct:.1f}%', format_amount(wa), note])
        add_styled_table(doc, ['阶段', '数量', '金额', '占比', '加权金额', '说明'], sd_rows, [2.5, 2, 3, 2, 3, 4])

        for cff in chart_files:
            if 'stage_distribution' in cff:
                doc.add_picture(cff, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                break
    doc.add_page_break()

    # ---- 第二章 ----
    h1 = doc.add_heading('二、漏斗健康度诊断', level=1)
    for run in h1.runs: set_run_font(run, font_name='宋体')
    h2 = doc.add_heading('2.1 漏斗形状诊断', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    si = []
    if stage_stats:
        ta = sum(v['amount'] for v in stage_stats.values())
        for sn in STAGE_ORDER:
            if sn in stage_stats:
                pct = stage_stats[sn]['amount'] / max(ta, 1) * 100
                if sn in ('客户立项', '技术认可') and pct > 25:
                    si.append(f'•「{sn}」占比{pct:.1f}%，项目堆积严重。')
                if sn == '签约准备' and pct < 5 and ta > 0:
                    si.append(f'•「签约准备」仅占{pct:.1f}%，短期订单不足。')
    if si:
        for issue in si:
            p = doc.add_paragraph()
            r = p.add_run(issue)
            set_run_font(r, font_name='宋体', size=Pt(10))
            if '⚠️' in issue: r.font.color.rgb = RGBColor(231, 76, 60)
    else:
        doc.add_paragraph('漏斗形状相对健康。')

    h2 = doc.add_heading('2.2 漏斗流速诊断', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    stag_count = {}
    for rec in analyzer.invalid_records:
        if 'stagnant' in rec['reason_code']:
            stag_count[rec.get('stage', '未知')] = stag_count.get(rec.get('stage', '未知'), 0) + 1
    if stag_count:
        sd2 = []
        for stg, cnt in sorted(stag_count.items(), key=lambda x: -x[1]):
            risk = '🔴 极高' if cnt > 5 else ('🟡 高' if cnt > 2 else '🟢 中')
            sd2.append([stg, str(cnt), risk, '核实推进状态' if cnt <= 5 else '立即清理'])
        add_styled_table(doc, ['停滞阶段', '数量', '风险', '建议'], sd2, [3, 3, 3, 7])

    for cff in chart_files:
        if 'invalid_reasons' in cff:
            doc.add_picture(cff, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            break
    doc.add_page_break()

    # ---- 第三章 ----
    h1 = doc.add_heading('三、无效业务机会识别与剔除', level=1)
    for run in h1.runs: set_run_font(run, font_name='宋体')
    h2 = doc.add_heading('3.1 无效机会判定标准', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    p = doc.add_paragraph('以下标准对每条记录逐一判定，命中任一即标记为"无效机会"。')
    for run in p.runs: set_run_font(run, font_name='宋体')
    rd = []
    for i, item in enumerate(INVALID_REASON_MAP.items(), 1):
        val_str = str(item[1]) if not isinstance(item[1], str) else item[1]
        parts = val_str.split('：', 1)
        rd.append([f'R{i}', parts[0] if parts else val_str, parts[1] if len(parts) > 1 else ''])
    add_styled_table(doc, ['规则', '判定标准', '说明'], rd, [2, 5, 9])

    h2 = doc.add_heading('3.2 无效机会明细', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    p = doc.add_paragraph(f'共 {len(analyzer.invalid_records)} 条无效记录：')
    for run in p.runs: set_run_font(run, font_name='宋体')
    rg = defaultdict(list)
    for rec in analyzer.invalid_records:
        rg[rec['reason_code']].append(rec)
    for code, records in sorted(rg.items(), key=lambda x: -len(x[1])):
        val_str = str(INVALID_REASON_MAP.get(code, code))
        short_title = val_str.split('：')[0] if val_str else str(code)
        h3 = doc.add_heading(f'▌{short_title}（{len(records)}条）', level=3)
        for run in h3.runs: set_run_font(run, font_name='宋体')
        p = doc.add_paragraph(INVALID_REASON_MAP.get(code, code))
        for run in p.runs: set_run_font(run, font_name='宋体')
        inv_d = []
        for rec in records[:50]:
            row = analyzer.df.iloc[rec['index']]
            inv_d.append([
                str(row.get('业务机会名', ''))[:35],
                str(rec['stage']), str(rec['owner']),
                format_amount(rec['amount']),
                format_date(row.get('预计签约日期', pd.NaT)),
                format_date(rec.get('stage_mod_date', pd.NaT)),
            ])
        add_styled_table(doc, ['业务机会名', '阶段', '负责人', '金额', '预计签约', '阶段修改'], inv_d, [5, 1.8, 1.5, 2, 2, 2])
        if len(records) > 50:
            p = doc.add_paragraph()
            r = p.add_run(f'... 还有{len(records) - 50}条省略')
            set_run_font(r, font_name='宋体', size=Pt(9), color=RGBColor(149, 165, 166))
    doc.add_page_break()

    # ---- 第四章 (合并有效业务机会甄别与高价值机会) ----
    h1 = doc.add_heading('四、有效业务机会甄别与高价值机会清单', level=1)
    for run in h1.runs: set_run_font(run, font_name='宋体')
    h2 = doc.add_heading('4.1 有效业务机会甄别过程', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    for cff in chart_files:
        if '有效业务机会_funnel' in cff:
            doc.add_picture(cff, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            break

    descs = [
        'S1 - 排除线索和公司立项(成交<10%)。',
        'S2 - 保留客户立项及以后。',
        'S3 - 排除不可控渠道(需接触决策人)。',
        'S4 - 120天内活跃更新。',
        'S5 - 金额>5000元。',
    ]
    for desc in descs:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(desc)
        set_run_font(r, font_name='宋体', size=Pt(10))

    final_key = '甄别S5｜有效金额>5000元'
    final_df = analyzer.有效业务机会_filters.get(final_key, pd.DataFrame())
    if not final_df.empty:
        p2 = doc.add_paragraph()
        r2 = p2.add_run(f'✅ 经过5轮甄别，最终保留 {len(final_df)} 个真机会，')
        set_run_font(r2, font_name='宋体', size=Pt(10), color=RGBColor(39, 174, 96))
        a_s = pd.to_numeric(safe_col(final_df, '预计签约金额', 0), errors='coerce').fillna(0)
        r3 = p2.add_run(f'金额 {format_amount(float(a_s.sum()))}')
        set_run_font(r3, font_name='宋体', size=Pt(10), color=RGBColor(39, 174, 96))

    h2 = doc.add_heading('4.2 高价值机会评分排序', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    p = doc.add_paragraph('对甄别后的真机会进行多维评分（满分100分：阶段50分+金额30分+直签20分），排序得出高价值清单：')
    for run in p.runs: set_run_font(run, font_name='宋体', size=Pt(10))

    if not analyzer.high_value.empty:
        for cff in chart_files:
            if 'high_value_top20' in cff:
                doc.add_picture(cff, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                break

        h2 = doc.add_heading('4.3 全部高价值机会', level=2)
        for run in h2.runs: set_run_font(run, font_name='宋体')
        hv_d = []
        for rank, (_, row) in enumerate(analyzer.high_value.iterrows(), 1):
            if '_is_direct' in row.index and row['_is_direct']:
                st = '⭐' + str(row.get('销售类型', ''))
            else:
                st = str(row.get('销售类型', ''))
            hv_d.append([
                str(rank), str(row.get('业务机会名', ''))[:30], str(row.get('阶段', '')),
                str(row.get('业务机会所有人', '')), st,
                format_amount(row.get('_amt', 0)), format_amount(row.get('_weighted_amt', 0)),
                f'{row.get("_total_score", 0):.0f}', format_date(row.get('创建日期', pd.NaT)),
                str(row.get('行业', '')),
            ])
        add_styled_table(doc, ['排名', '业务机会名', '阶段', '负责人', '销售类型', '金额', '加权', '评分', '创建', '行业'], hv_d[:80], [1, 4.5, 1.5, 1.5, 1.5, 1.8, 1.8, 0.8, 1.8, 2])
        if len(hv_d) > 80:
            p = doc.add_paragraph(f'... 共{len(hv_d)}条，仅显示前80条。')
            for run in p.runs: set_run_font(run, font_name='宋体')

        h2 = doc.add_heading('4.4 Top 10 详情', level=2)
        for run in h2.runs: set_run_font(run, font_name='宋体')
        for rank, (_, row) in enumerate(analyzer.high_value.head(10).iterrows(), 1):
            h3 = doc.add_heading(f'#{rank} {str(row.get("业务机会名", ""))[:50]}', level=3)
            for run in h3.runs: set_run_font(run, font_name='宋体')
            fields = [
                ('业务机会编号', row.get('业务机会编号', '')), ('阶段', row.get('阶段', '')),
                ('可能性', row.get('可能性', '')), ('负责人', row.get('业务机会所有人', '')),
                ('预计签约金额', format_amount(row.get('_amt', 0))),
                ('加权金额', format_amount(row.get('_weighted_amt', 0))),
                ('评分', f'{row.get("_total_score", 0):.0f}分'),
                ('预计签约日期', format_date(row.get('预计签约日期', pd.NaT))),
                ('阶段修改日期', format_date(row.get('阶段修改日期', pd.NaT))),
                ('创建日期', format_date(row.get('创建日期', pd.NaT))),
                ('最近更新', format_date(row.get('更新时间', pd.NaT))),
                ('客户名称', str(row.get('客户名', ''))[:40]), ('行业', row.get('行业', '')),
                ('销售类型', row.get('销售类型', '')), ('销售部门', row.get('销售部门', '')),
            ]
            dd = []
            for f_name, f_val in fields:
                if pd.notna(f_val) and str(f_val).strip() and str(f_val) != 'nan':
                    dd.append([f_name, str(f_val)[:50]])
            add_styled_table(doc, ['字段', '内容'], dd, [3, 13])
            doc.add_paragraph('')
    else:
        p = doc.add_paragraph('⚠️ 未识别出高价值机会。')
        for run in p.runs: set_run_font(run, font_name='宋体')
    doc.add_page_break()

    # ---- 第五章 (原第六章) ----
    h1 = doc.add_heading('五、人员漏斗透视', level=1)
    for run in h1.runs: set_run_font(run, font_name='宋体')
    for cff in chart_files:
        if 'owner_comparison' in cff:
            doc.add_picture(cff, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            break
    ow = analyzer.summary.get('owner_distribution', {})
    if ow:
        od_rows = []
        for n in sorted(ow.keys(), key=lambda x: -ow[x]['amount']):
            v = ow[n]
            ratio = v['weighted'] / v['amount'] * 100 if v['amount'] > 0 else 0
            if ratio < 25: risk = '🔴 漏斗虚胖'
            elif ratio < 40: risk = '🟡 需关注'
            else: risk = '🟢 健康'
            od_rows.append([n, str(v['count']), format_amount(v['amount']), format_amount(v['weighted']), f'{ratio:.1f}%', risk])
        add_styled_table(doc, ['负责人', '有效机会', '名义金额', '加权金额', '折算率', '风险提示'], od_rows, [2, 2, 3, 3, 2, 6])
    doc.add_page_break()

    # ---- 第六章 (原第七章) ----
    h1 = doc.add_heading('六、管理行动建议', level=1)
    for run in h1.runs: set_run_font(run, font_name='宋体')
    h2 = doc.add_heading('6.1 漏斗清理', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    for a in [
        f'1. 清理{len(analyzer.invalid_records)}条无效记录。',
        '2. 远古僵尸单>3年逐项提供推进计划。',
        '3. 签约准备停滞(阶段修改>180天)补充签报号或降级。',
        '4. 纯借机单独管理。',
    ]:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(a)
        set_run_font(r, font_name='宋体', size=Pt(10))

    h2 = doc.add_heading('6.2 高价值保护', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    for a in [
        f'1. 跟踪{analyzer.summary["high_value_count"]}个高价值机会' f'(加权{format_amount(analyzer.summary["high_value_weighted"])})。',
        '2. Top10指派售前资源。',
        '3. 周会有效业务机会检查。',
        '4. ⭐直签项目优先。',
    ]:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(a)
        set_run_font(r, font_name='宋体', size=Pt(10))

    h2 = doc.add_heading('6.3 阶段准入强化', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    for a in [
        '1. 签约准备须有签报号+回款计划+阶段修改日期在90天内。',
        '2. 商务认可须有沟通记录+阶段修改日期在180天内。',
        '3. 技术认可须有技术方案（POC）确认+阶段修改日期在180天内。',
        '4. 超期未推进自动降级。',
    ]:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(a)
        set_run_font(r, font_name='宋体', size=Pt(10))

    doc.add_paragraph('')
    pe = doc.add_paragraph(); pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pe.add_run('—— 报告结束 ——')
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(149, 165, 166)

    doc.add_page_break()
    doc.add_heading('附录：高价值机会评分标准', level=1)
    doc.add_paragraph('高价值机会评分由三个维度构成，满分100分，具体标准如下：')
    doc.add_heading('1. 阶段得分（满分50分）', level=2)
    doc.add_paragraph('根据业务机会当前所处销售阶段，乘以对应权重系数再乘以50：')
    stage_score_rows = [
        ['销售线索', '5%', '2.5分'], ['公司立项', '10%', '5分'], ['客户立项', '30%', '15分'],
        ['技术认可', '50%', '25分'], ['商务认可', '75%', '37.5分'], ['签约准备', '100%', '50分'],
        ['已赢单', '100%', '50分'], ['已输单', '0%', '0分'],
    ]
    add_styled_table(doc, ['阶段', '权重', '得分'], stage_score_rows, [4, 4, 4])

    doc.add_heading('2. 金额得分（满分30分）', level=2)
    doc.add_paragraph('根据预计签约金额的大小划分为四档：')
    amount_score_rows = [
        ['50万元及以上', '30分'], ['10万元（含）至50万元', '20分'],
        ['3万元（含）至10万元', '10分'], ['5000元以上至3万元', '5分'],
    ]
    add_styled_table(doc, ['金额区间', '得分'], amount_score_rows, [6, 4])

    doc.add_heading('3. 销售类型加分（满分20分）', level=2)
    type_score_rows = [ ['直签用户', '直接加20分'], ['其他类型（如渠道销售）', '0分'], ]
    add_styled_table(doc, ['销售类型', '得分'], type_score_rows, [6, 4])

    doc.add_paragraph('')
    p_note = doc.add_paragraph()
    r_note = p_note.add_run('综合评分 = 阶段得分 + 金额得分 + 销售类型加分')
    r_note.font.size = Pt(10); r_note.bold = True; r_note.font.color.rgb = RGBColor(44, 62, 80)

    # ==========================================
    # 彻底清除因图片插入产生的空段落空白页
    # ==========================================
    remove_empty_paragraph_before_picture(doc)

    # ==========================================
    # 仅增加页码
    # ==========================================
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)
    run4 = p.add_run("1")
    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)
    for r in [run, run2, run3, run4, run5]:
        r.font.name = '宋体'
        r.font.size = Pt(9)

    doc.save(output_path)
    return output_path

# ================================================================
# UI界面
# ================================================================
class FunnelAnalysisApp:
    def __init__(self, root):
        self.root = root
        self.root.title('销售漏斗智能分析工具 v1.0')
        self.root.geometry('750x580')
        self.root.minsize(600, 400)
        self.root.update_idletasks()
        x = (self.root.winfo_screenwidth() - 750) // 2
        y = (self.root.winfo_screenheight() - 580) // 2
        self.root.geometry(f'750x580+{x}+{y}')
        self.file_path = tk.StringVar()
        self.output_dir = tk.StringVar()
        self.analyzer = None
        self._build_ui()

    def _build_ui(self):
        bg = '#ECF0F1'; hc = '#2C3E50'; ac = '#3498DB'
        self.root.configure(bg=bg)
        hf = tk.Frame(self.root, bg=hc, height=60)
        hf.pack(fill='x'); hf.pack_propagate(False)
        tk.Label(hf, text='📊 销售漏斗智能分析工具 v1.0', font=('微软雅黑', 16, 'bold'), bg=hc, fg='white').pack(side='left', padx=20, pady=15)
        tk.Label(hf, text='筛选2026年签约 ', font=('微软雅黑', 9), bg=hc, fg='#BDC3C7').pack(side='right', padx=20)

        cf = tk.Frame(self.root, bg=bg)
        cf.pack(fill='both', expand=True, padx=20, pady=15)

        ff = tk.LabelFrame(cf, text=' 📁 数据源 ', font=('微软雅黑', 10, 'bold'), bg=bg, fg=hc, padx=15, pady=10)
        ff.pack(fill='x', pady=(0, 10))
        tk.Label(ff, text='业务机会Excel文件：', font=('微软雅黑', 9), bg=bg).pack(anchor='w')
        r1 = tk.Frame(ff, bg=bg); r1.pack(fill='x', pady=3)
        tk.Entry(r1, textvariable=self.file_path, font=('微软雅黑', 9), state='readonly', bg='white').pack(side='left', fill='x', expand=True, padx=(0, 10))
        tk.Button(r1, text='选择文件', command=self._select_file, font=('微软雅黑', 9), bg=ac, fg='white', relief='flat', padx=15).pack(side='right')

        of = tk.LabelFrame(cf, text=' 💾 输出设置 ', font=('微软雅黑', 10, 'bold'), bg=bg, fg=hc, padx=15, pady=10)
        of.pack(fill='x', pady=(0, 10))
        tk.Label(of, text='报告保存路径：', font=('微软雅黑', 9), bg=bg).pack(anchor='w')
        r2 = tk.Frame(of, bg=bg); r2.pack(fill='x', pady=3)
        tk.Entry(r2, textvariable=self.output_dir, font=('微软雅黑', 9), state='readonly', bg='white').pack(side='left', fill='x', expand=True, padx=(0, 10))
        tk.Button(r2, text='选择路径', command=self._select_output, font=('微软雅黑', 9), bg=ac, fg='white', relief='flat', padx=15).pack(side='right')

        inf = tk.LabelFrame(cf, text=' 📋 分析内容 ', font=('微软雅黑', 10, 'bold'), bg=bg, fg=hc, padx=15, pady=8)
        inf.pack(fill='x', pady=(0, 10))
        tk.Label(inf, text='✅ 仅分析2026年预计签约项目 ✅ 签约准备停滞用阶段修改日期\n' '✅ 技术认可/商务认可停滞用阶段修改日期 ✅ 有效业务机会五步甄别\n' '✅ 高价值机会清单 ✅ 人员漏斗透视 ✅ 管理行动建议', font=('微软雅黑', 8), bg=bg, fg='#7F8C8D', justify='left').pack(anchor='w')

        bf = tk.Frame(cf, bg=bg); bf.pack(fill='x', pady=(0, 10))
        self.run_btn = tk.Button(bf, text='🚀 开始分析并生成报告', command=self._run, font=('微软雅黑', 11, 'bold'), bg='#27AE60', fg='white', relief='flat', padx=30, pady=8)
        self.run_btn.pack()

        lf = tk.LabelFrame(cf, text=' 📝 运行日志 ', font=('微软雅黑', 10, 'bold'), bg=bg, fg=hc, padx=10, pady=8)
        lf.pack(fill='both', expand=True)
        self.log_w = tk.Text(lf, font=('Consolas', 9), bg='#2C3E50', fg='#2ECC71', height=10, wrap='word', relief='flat')
        self.log_w.pack(fill='both', expand=True)
        sb = tk.Scrollbar(self.log_w, command=self.log_w.yview)
        sb.pack(side='right', fill='y')
        self.log_w.config(yscrollcommand=sb.set)

        self.status = tk.StringVar(value='就绪')
        tk.Label(self.root, textvariable=self.status, font=('微软雅黑', 8), bg='#2C3E50', fg='#BDC3C7', anchor='w', padx=10, pady=3).pack(fill='x', side='bottom')

    def _log(self, msg):
        ts = datetime.now().strftime('%H:%M:%S')
        self.log_w.insert('end', f'[{ts}] {msg}\n')
        self.log_w.see('end')
        self.root.update_idletasks()

    def _select_file(self):
        p = filedialog.askopenfilename( title='选择业务机会Excel', filetypes=[('Excel', '*.xlsx *.xls *.csv'), ('所有文件', '*.*')])
        if p:
            self.file_path.set(p)
            self.status.set(f'已选择：{os.path.basename(p)}')

    def _select_output(self):
        p = filedialog.askdirectory(title='选择保存路径')
        if p:
            self.output_dir.set(p)
            self.status.set(f'输出：{p}')

    def _run(self):
        fp = self.file_path.get().strip()
        od = self.output_dir.get().strip()
        if not fp: messagebox.showwarning('提示', '请选择Excel文件！'); return
        if not os.path.exists(fp): messagebox.showerror('错误', '文件不存在！'); return
        if not od: od = os.path.dirname(fp); self.output_dir.set(od)
        self.run_btn.config(state='disabled', text='⏳ 分析中...')
        self.status.set('正在分析...')
        self.log_w.delete('1.0', 'end')
        try:
            self._do(fp, od)
        except Exception as e:
            self._log(f'❌ 失败：{str(e)}')
            messagebox.showerror('错误', f'分析失败：\n{str(e)}')
            import traceback; traceback.print_exc()
        finally:
            self.run_btn.config(state='normal', text='🚀 开始分析并生成报告')

    def _do(self, fp, od):
        self._log('📖 纯pandas读取Excel...')
        try:
            df, header_idx = smart_read_excel(fp)
        except Exception as e:
            self._log(f'❌ 读取失败: {e}'); raise e
        self._log(f' ✅ 表头行: 第{header_idx + 1}行')
        self._log(f' ✅ {len(df)}条记录, {len(df.columns)}个字段')
        cols_display = list(df.columns[:8])
        if len(df.columns) > 8: cols_display.append('...')
        self._log(f' 📋 列名: {", ".join(cols_display)}')
        found = [c for c in REQUIRED_COLS if c in df.columns]
        missing = [c for c in REQUIRED_COLS if c not in df.columns]
        if found: self._log(f' ✅ 已识别({len(found)}/{len(REQUIRED_COLS)}): ' f'{", ".join(found)}')
        if missing: self._log(f' ⚠️ 缺失({len(missing)}/{len(REQUIRED_COLS)}): ' f'{", ".join(missing)}')
        self._log('🔍 列名增强映射...')
        cm = map_columns_enhanced(df)
        if cm:
            for old, new in cm.items(): df[new] = df[old]
            self._log(f' ✅ 映射: {", ".join(f"{k}→{v}" for k, v in cm.items())}')
        else:
            self._log(' ✅ 列名已匹配')
        found2 = [c for c in REQUIRED_COLS if c in df.columns]
        self._log(f' 📋 最终识别: {len(found2)}/{len(REQUIRED_COLS)}')
        for c in ['阶段', '预计签约金额', '业务机会所有人', '业务机会名', '预计签约日期', '阶段修改日期']:
            if c in df.columns:
                sample = df[c].dropna().head(3).tolist()
                self._log(f' 📋 [{c}] {sample}')
        self._log('⚙️ 第0步：筛选2026年签约项目...')
        self._log('⚙️ 运行分析引擎...')
        self.analyzer = FunnelAnalyzer(df)
        summary = self.analyzer.run_analysis()
        yf = summary.get('read_diagnosis', {}).get('year_filter', {})
        if yf:
            self._log(f' 📅 原始{yf.get("before",0)}条 → ' f'排除{yf.get("excluded",0)}条非2026年 → ' f'保留{yf.get("after",0)}条')
        self._log(f' 总:{summary["total_count"]} 无效:{summary["invalid_count"]}' f' 有效:{summary["valid_count"]} 高价值:{summary["high_value_count"]}')
        self._log(f' 总金额:{format_amount(summary["total_amount"])}' f' 有效金额:{format_amount(summary["valid_amount"])}' f' 加权金额:{format_amount(summary["high_value_weighted"])}')
        if self.analyzer.missing_cols:
            self._log(f' ⚠️ 仍有缺失列: {", ".join(self.analyzer.missing_cols)}')
        self._log('📊 生成图表...')
        cd = os.path.join(od, 'charts')
        charts = generate_charts(self.analyzer, cd)
        self._log(f' ✅ {len(charts)}张图表')
        self._log('📝 生成DOCX报告...')
        ts = datetime.now().strftime('%Y%m%d_%H%M')
        rn = f'销售漏斗分析报告_{ts}.docx'
        rp = os.path.join(od, rn)
        generate_report(self.analyzer, charts, rp)
        self._log(f' ✅ 已保存: {rp}')
        self.status.set(f'✅ 完成: {rn}')
        self._log('🎉 分析完成！')
        messagebox.showinfo('完成', f'报告已生成：\n{rp}')

# ================================================================
# 入口
# ================================================================
if __name__ == '__main__':
    root = tk.Tk()
    FunnelAnalysisApp(root)
    root.mainloop()
