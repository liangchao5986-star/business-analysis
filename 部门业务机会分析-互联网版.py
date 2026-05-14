""" 销售漏斗智能分析工具 v1.0 """
import os, re, tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
import numpy as np
from datetime import datetime, timedelta
from collections import defaultdict
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ================================================================
# 常量
# ================================================================
STAGE_ORDER = ['销售线索', '公司立项', '客户立项', '技术认可', '商务认可', '签约准备', '已赢单', '已输单']
STAGE_WEIGHT = {
    '销售线索': 0.05, '公司立项': 0.10, '客户立项': 0.30, '技术认可': 0.50,
    '商务认可': 0.75, '签约准备': 1.00, '已赢单': 1.00, '已输单': 0.00
}
STAGE_WEIGHT_ORDERED = {s: STAGE_WEIGHT.get(s, 0.1) for s in STAGE_ORDER}
INVALID_REASON_MAP = {
    'ancient': '🗑️ 远古僵尸单：创建超过3年仍未签约',
    'not_2026': '🗑️ 非本年度项目：预计签约日期不在2026年',
    'stagnant_prep': '🗑️ 签约准备停滞：阶段修改日期距今>180天',
    'stagnant_tech': '🗑️ 技术认可停滞：阶段修改日期距今>360天',
    'stagnant_comm': '🗑️ 商务认可停滞：阶段修改日期距今>360天',
    'zero_amount': '🗑️ 金额为零：无法带来收入贡献',
    'zero_prob': '🗑️ 阶段概率不匹配：后期阶段可能性过低',
    'no_update': '🗑️ 超180天未更新：疑似已放弃跟进',
    'pure_borrow': '🗑️ 纯借机项目：无直接销售收入',
    'tiny_amount': '⚠️ 微型金额(<5000元)',
}
REQUIRED_COLS = ['阶段', '创建日期', '预计签约金额', '业务机会所有人', '销售类型', '更新时间', '业务机会名', '项目类型', '可能性', '预计签约日期', '阶段修改日期']
ALL_STANDARD_NAMES = [
    '阶段', '预计签约金额', '业务机会所有人', '业务机会名', '创建日期', '销售类型', '更新时间', '项目类型', '可能性', '行业', '客户名', '销售部门', '产品描述', '预计签约日期', '未签报预计回款金额', '业务机会编号', '结果', '主要市场活动来源', '是否Name客户', '是否有定制开发', '已转签报单', '签约客户联系人', '最终客户', '最终客户联系人', '丢单原因', '竞争对手', '丢单原因备注', '价格手册', '客户编号', '预计供货日期', '是否预计签约当月验收', '阶段修改日期', '项目名称', '项目编号', '签报单流水号', '合同名称', '合同编号', '合计', '所有人编号', '所属公司', '原销售人员', '所有人部门', '创建人', '创建时间', '修改时间', '事业部',
]

# ================================================================
# 工具函数
# ================================================================
def safe_col(df, col_name, fallback=''):
    if col_name in df.columns:
        return df[col_name]
    return pd.Series([fallback] * len(df), index=df.index, name=col_name)

def format_amount(val):
    try:
        v = float(val)
        return f'{v / 10000:.2f}万' if abs(v) >= 10000 else f'{v:,.0f}元'
    except:
        return str(val) if val else '-'

def format_date(val):
    try:
        if pd.isna(val): return '-'
        if isinstance(val, (pd.Timestamp, datetime)): return val.strftime('%Y-%m-%d')
        return str(val)[:10]
    except:
        return '-'

def safe_timedelta(dt):
    if pd.isna(dt): return 0
    try:
        return (datetime.now() - pd.to_datetime(dt)).days
    except:
        return 0

def clean_col_name(name):
    s = str(name)
    s = re.sub(r'[\u200b\u200c\u200d\ufeff\u00a0\u200e\u200f]', '', s)
    s = s.strip()
    s = re.sub(r'\s+', ' ', s)
    return s

# ================================================================
# 智能Excel读取（纯 pandas，杜绝 ZIP closed）
# ================================================================
def smart_read_excel(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.csv':
        df = pd.read_csv(file_path)
        df.columns = [clean_col_name(c) for c in df.columns]
        return df, 0
    best_df = None
    best_header = 0
    best_score = 0
    for h in range(6):
        try:
            test_df = pd.read_excel(file_path, header=h)
            test_df.columns = [clean_col_name(c) for c in test_df.columns]
            score = sum(1 for c in test_df.columns if c in ALL_STANDARD_NAMES)
            if score > best_score:
                best_score = score
                best_header = h
                best_df = test_df
        except Exception:
            continue
    if best_df is not None and best_score >= 3:
        best_df.columns = [clean_col_name(c) for c in best_df.columns]
        return best_df, best_header
    try:
        df = pd.read_excel(file_path)
        df.columns = [clean_col_name(c) for c in df.columns]
        return df, 0
    except Exception:
        raise ValueError(f'无法读取Excel文件：{file_path}')

# ================================================================
# 列名映射
# ================================================================
def map_columns_enhanced(df):
    col_map = {}
    cols = list(df.columns)
    rules = {
        '业务机会名': ['业务机会名', '商机名称', '商机名', '项目名称', '项目名', '名称', '机会名称'],
        '创建日期': ['创建日期', '创建时间', '录入日期', '录入时间', '创建日', '建档日期'],
        '阶段': ['阶段', '商机阶段', '业务阶段', '当前阶段', '阶段名称', '所处阶段', '状态'],
        '预计签约金额': ['预计签约金额', '预计金额', '金额', '预计签约额', '签约金额', '预估金额'],
        '业务机会所有人': ['业务机会所有人', '负责人', '销售', '销售负责人', '所有人', '销售员', '客户经理', '销售人'],
        '业务机会编号': ['业务机会编号', '商机编号', '编号', 'OPP编号', '机会编号'],
        '行业': ['行业', '行业分类', '行业类别', '客户行业'],
        '销售类型': ['销售类型', '销售模式', '签约方式', '类型'],
        '更新时间': ['更新时间', '修改时间', '最后更新', '更新日期', '最后修改', '修改日期'],
        '预计签约日期': ['预计签约日期', '预计签约', '计划签约日期', '签约日期', '计划签约'],
        '未签报预计回款金额': ['未签报预计回款金额', '预计回款金额', '回款金额', '预计回款'],
        '项目类型': ['项目类型', '项目类别', '商机类型'],
        '客户名': ['客户名', '客户名称', '客户', '客户简称'],
        '销售部门': ['销售部门', '部门', '所属部门', '团队'],
        '产品描述': ['产品描述', '产品', '产品名称', '产品线'],
        '可能性': ['可能性', '赢率', '概率', '胜率', '成交概率', '预计赢率'],
        '阶段修改日期': ['阶段修改日期', '阶段变更日期', '阶段更新日期', '阶段推进日期'],
    }
    for standard, aliases in rules.items():
        if standard in cols: continue
        matched = False
        for alias in aliases:
            alias_str = str(alias) if not isinstance(alias, str) else alias
            for c in cols:
                if c.lower() == alias_str.lower():
                    col_map[c] = standard; matched = True; break
            if matched: break
        if not matched:
            for alias in aliases:
                alias_str = str(alias) if not isinstance(alias, str) else alias
                if len(alias_str) < 2: continue
                for c in cols:
                    if alias_str in c or c in alias_str:
                        col_map[c] = standard; matched = True; break
                if matched: break
        if not matched:
            for c in cols:
                if len(standard) >= 2 and standard in c:
                    col_map[c] = standard; break
    return col_map

# ================================================================
# 分析引擎
# ================================================================
class FunnelAnalyzer:
    def __init__(self, df):
        self.df = df.copy()
        self.df.columns = [clean_col_name(c) for c in self.df.columns]
        self._parse_dates()
        self._parse_amounts()
        self.invalid_records = []
        self.有效业务机会_filters = {}
        self.high_value = pd.DataFrame()
        self.all_valid = pd.DataFrame()
        self.summary = {}
        self.missing_cols = []
        self.read_diagnosis = {}
        self.excluded_by_year = 0

    def _parse_dates(self):
        for c in self.df.columns:
            for kw in ['日期', '时间', 'date', 'time']:
                if kw in c.lower():
                    self.df[c] = pd.to_datetime(self.df[c], errors='coerce')
                    break

    def _parse_amounts(self):
        for c in self.df.columns:
            for kw in ['金额', '回款', 'amount', '合计']:
                if kw in c.lower():
                    self.df[c] = pd.to_numeric(self.df[c], errors='coerce').fillna(0)
                    break

    def _ensure_col(self, col_name):
        if col_name not in self.df.columns:
            if col_name not in self.missing_cols:
                self.missing_cols.append(col_name)
            self.df[col_name] = ''

    def run_analysis(self):
        self.read_diagnosis = {
            'total_columns': len(self.df.columns),
            'column_names': list(self.df.columns),
        }
        for col in REQUIRED_COLS:
            self._ensure_col(col)
        self._step0_filter_year()
        self._step1()
        self._step2()
        self._step3()
        self._step4()
        self._step5()
        return self.summary

    def _step0_filter_year(self):
        """第0步：筛选预计签约日期为2026年的项目"""
        col = '预计签约日期'
        total_before = len(self.df)
        if col in self.df.columns:
            mask = self.df[col].notna()
            year_mask = pd.Series([False] * len(self.df), index=self.df.index)
            for idx in self.df.index:
                val = self.df.at[idx, col]
                if pd.notna(val):
                    try:
                        dt = pd.to_datetime(val)
                        if dt.year == 2026:
                            year_mask.at[idx] = True
                    except: pass
            excluded = mask & ~year_mask
            self.excluded_by_year = int(excluded.sum())
            self.df = self.df[year_mask | ~mask].copy().reset_index(drop=True)
        else:
            self.excluded_by_year = 0
        after = len(self.df)
        self.read_diagnosis['year_filter'] = {
            'before': total_before, 'excluded': self.excluded_by_year, 'after': after,
        }

    def _step1(self):
        self.invalid_records = []
        self.df['_invalid_reasons'] = [[] for _ in range(len(self.df))]
        for i, row in self.df.iterrows():
            reasons = []
            create_dt = row.get('创建日期', pd.NaT)
            update_dt = row.get('更新时间', pd.NaT)
            stage_mod_dt = row.get('阶段修改日期', pd.NaT)
            stage = str(row.get('阶段', '') or '').strip()
            amt = float(row.get('预计签约金额', 0) or 0)
            proj_type = str(row.get('项目类型', '') or '')
            prob_str = str(row.get('可能性', '') or '')
            age = safe_timedelta(create_dt)
            if age > 365 * 3: reasons.append('ancient')
            stag_age = safe_timedelta(stage_mod_dt)
            if stag_age == 0: stag_age = safe_timedelta(update_dt)
            if stag_age == 0: stag_age = age
            if stage == '签约准备' and stag_age > 180: reasons.append('stagnant_prep')
            elif stage == '技术认可' and stag_age > 360: reasons.append('stagnant_tech')
            elif stage == '商务认可' and stag_age > 360: reasons.append('stagnant_comm')
            if amt == 0: reasons.append('zero_amount')
            if stage in ('商务认可', '签约准备'):
                try:
                    p = float(prob_str.replace('%', '')) if '%' in prob_str else float(prob_str)
                    if stage == '签约准备' and p < 90: reasons.append('zero_prob')
                except: pass
            if safe_timedelta(update_dt) > 180: reasons.append('no_update')
            if '纯借机' in proj_type: reasons.append('pure_borrow')
            if 0 < amt < 5000: reasons.append('tiny_amount')
            self.df.at[i, '_invalid_reasons'] = reasons
            for r in reasons:
                self.invalid_records.append({
                    'index': i, 'name': str(row.get('业务机会名', f'行{i}') or '')[:50],
                    'stage': stage, 'owner': str(row.get('业务机会所有人', '') or ''),
                    'amount': amt, 'create_date': create_dt, 'stage_mod_date': stage_mod_dt,
                    'reason_code': r, 'reason_text': INVALID_REASON_MAP.get(r, r),
                })
        self.df['_is_invalid'] = self.df['_invalid_reasons'].apply(lambda x: len(x) > 0)

    def _step2(self):
        self.all_valid = self.df[~self.df['_is_invalid']].copy()

    def _step3(self):
        df = self.all_valid.copy()
        steps = {'原始有效机会': df.copy()}
        s_s = safe_col(df, '阶段', '')
        s1 = df[~s_s.isin(['销售线索', '公司立项'])].copy()
        steps['甄别S1｜排除销售线索与公司立项'] = s1
        s1s = safe_col(s1, '阶段', '')
        s2 = s1[s1s.isin(['客户立项', '技术认可', '商务认可', '签约准备'])].copy()
        steps['甄别S2｜保留客户立项及以后'] = s2
        st = safe_col(s2, '销售类型', '')
        m3 = st.isin(['直签用户', '渠道销售', ''])
        s3 = s2[m3].copy()
        s3['_is_direct'] = (st[m3] == '直签用户')
        steps['甄别S3｜排除不可控渠道类型'] = s3
        uc = safe_col(s3, '更新时间', pd.NaT)
        if uc.notna().any():
            cutoff = datetime.now() - timedelta(days=120)
            try:
                m4 = uc.notna() & (pd.to_datetime(uc, errors='coerce') >= cutoff)
            except:
                m4 = pd.Series([True] * len(s3), index=s3.index)
            s4 = s3[m4].copy()
        else:
            s4 = s3.copy()
        steps['甄别S4｜120天内有活跃更新'] = s4
        ac = pd.to_numeric(safe_col(s4, '预计签约金额', 0), errors='coerce').fillna(0)
        s5 = s4[ac > 5000].copy()
        steps['甄别S5｜有效金额>5000元'] = s5
        self.有效业务机会_filters = steps

    def _step4(self):
        key = '甄别S5｜有效金额>5000元'
        if key not in self.有效业务机会_filters:
            self.high_value = pd.DataFrame(); return
        df = self.有效业务机会_filters[key].copy()
        if df.empty: self.high_value = pd.DataFrame(); return
        ss = safe_col(df, '阶段', '')
        an = pd.to_numeric(safe_col(df, '预计签约金额', 0), errors='coerce').fillna(0)
        df['_weight'] = ss.map(STAGE_WEIGHT_ORDERED).fillna(0.1)
        df['_amt'] = an
        df['_weighted_amt'] = an * df['_weight']
        df['_score_stage'] = df['_weight'] * 50
        df['_score_amount'] = 0
        df.loc[an >= 500000, '_score_amount'] = 30
        df.loc[(an >= 100000) & (an < 500000), '_score_amount'] = 20
        df.loc[(an >= 30000) & (an < 100000), '_score_amount'] = 10
        df.loc[(an > 5000) & (an < 30000), '_score_amount'] = 5
        df['_score_direct'] = 0
        if '_is_direct' in df.columns:
            df.loc[df['_is_direct'], '_score_direct'] = 20
        df['_total_score'] = df['_score_stage'] + df['_score_amount'] + df['_score_direct']
        self.high_value = df.sort_values( ['_total_score', '_weighted_amt'], ascending=[False, False] ).reset_index(drop=True)

    def _step5(self):
        total = len(self.df)
        ic = int(self.df['_is_invalid'].sum()) if '_is_invalid' in self.df.columns else 0
        vc = total - ic
        aa = pd.to_numeric(safe_col(self.df, '预计签约金额', 0), errors='coerce').fillna(0)
        ta = float(aa.sum())
        ia = float(aa[self.df['_is_invalid']].sum()) if ic > 0 else 0
        va = ta - ia
        hvc = len(self.high_value)
        hva = float(self.high_value['_amt'].sum()) if not self.high_value.empty else 0
        hvw = float(self.high_value['_weighted_amt'].sum()) if not self.high_value.empty else 0
        sd = {}
        ss = safe_col(self.df, '阶段', '')
        for s in STAGE_ORDER:
            m = ss == s; c = int(m.sum())
            if c > 0:
                sd[s] = {'count': c, 'amount': float(aa[m].sum())}
        od = {}
        if '业务机会所有人' in self.all_valid.columns and not self.all_valid.empty:
            for n, g in self.all_valid.groupby('业务机会所有人'):
                a = pd.to_numeric(safe_col(g, '预计签约金额', 0), errors='coerce').fillna(0)
                w = safe_col(g, '阶段', '').map(STAGE_WEIGHT_ORDERED).fillna(0.1)
                od[str(n)] = {
                    'count': len(g), 'amount': float(a.sum()), 'weighted': float((a * w).sum()),
                }
        self.summary = {
            'total_count': total, 'invalid_count': ic, 'valid_count': vc,
            'total_amount': ta, 'invalid_amount': ia, 'valid_amount': va,
            'high_value_count': hvc, 'high_value_amount': hva, 'high_value_weighted': hvw,
            'stage_distribution': sd, 'owner_distribution': od,
            'missing_cols': self.missing_cols, 'read_diagnosis': self.read_diagnosis,
            'excluded_by_year': self.excluded_by_year,
        }

# ================================================================
# 图表生成（解决字体方格问题，生成后嵌入文档不单独保存）
# ================================================================
def generate_charts(analyzer):
    import matplotlib
    # 解决matplotlib中文显示方格问题
    matplotlib.rcParams['font.family'] = ['SimHei', 'Microsoft YaHei', 'Arial Unicode MS', 'WenQuanYi Micro Hei', 'DejaVu Sans']
    matplotlib.rcParams['axes.unicode_minus'] = False
    matplotlib.use('Agg')  # 非交互式后端
    import matplotlib.pyplot as plt
    import io
    import base64

    charts = {}

    # 1. 销售漏斗各阶段金额分布
    ss = analyzer.summary.get('stage_distribution', {})
    if ss:
        fig, ax = plt.subplots(figsize=(10, 5))
        stages = [s for s in STAGE_ORDER if s in ss]
        amounts = [ss[s]['amount'] / 10000 for s in stages]
        counts = [ss[s]['count'] for s in stages]
        colors_list = ['#95a5a6', '#bdc3c7', '#3498db', '#2ecc71', '#f39c12', '#e74c3c']
        bar_colors = [colors_list[min(i, len(colors_list)-1)] for i in range(len(stages))]

        bars = ax.barh(range(len(stages)), amounts, color=bar_colors, edgecolor='white', height=0.6)
        for b, c, a in zip(bars, counts, amounts):
            ax.text(b.get_width() + max(amounts) * 0.01, b.get_y() + b.get_height() / 2,
                    f'{a:.1f}万({c}单)', va='center', fontsize=9)

        ax.set_yticks(range(len(stages)))
        ax.set_yticklabels(stages, fontsize=10)
        ax.set_xlabel('金额(万元)')
        ax.set_title('销售漏斗各阶段金额分布', fontsize=13, fontweight='bold')
        ax.invert_yaxis()
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        plt.tight_layout()

        # 保存到内存
        buf = io.BytesIO()
        plt.savefig(buf, dpi=150, bbox_inches='tight', format='png')
        buf.seek(0)
        charts['stage_distribution'] = buf
        plt.close()

    # 2. 无效机会原因分类
    if analyzer.invalid_records:
        fig, ax = plt.subplots(figsize=(9, 6))
        rc = {}
        for r in analyzer.invalid_records:
            rc[r['reason_code']] = rc.get(r['reason_code'], 0) + 1

        labels = []
        for k in rc.keys():
            val = str(INVALID_REASON_MAP.get(k, k))
            parts = val.split('：', 1)
            labels.append(parts[0] if parts else val)

        ax.pie(list(rc.values()), labels=labels, autopct='%1.1f%%',
               explode=[0.05]*len(rc), startangle=90, textprops={'fontsize':8}, pctdistance=0.85)
        ax.set_title(f'无效机会原因分类(共{len(analyzer.invalid_records)}条)', fontsize=13, fontweight='bold')
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, dpi=150, bbox_inches='tight', format='png')
        buf.seek(0)
        charts['invalid_reasons'] = buf
        plt.close()

    # 3. 销售人员有效漏斗对比
    od = analyzer.summary.get('owner_distribution', {})
    if od:
        fig, ax = plt.subplots(figsize=(12, 6))
        so = sorted(od.items(), key=lambda x: -x[1]['amount'])[:15]
        names = [o[0] for o in so]
        nom = [o[1]['amount'] / 10000 for o in so]
        wgt = [o[1]['weighted'] / 10000 for o in so]
        x = range(len(names))
        bw = 0.35

        ax.bar([i - bw/2 for i in x], nom, bw, label='名义金额', color='#3498db', alpha=0.8)
        ax.bar([i + bw/2 for i in x], wgt, bw, label='加权金额', color='#e74c3c', alpha=0.8)
        ax.set_xticks(x)
        ax.set_xticklabels(names, rotation=45, ha='right', fontsize=10)
        ax.set_ylabel('金额(万元)', fontsize=12)
        ax.set_title('销售人员有效漏斗对比', fontsize=14, fontweight='bold')
        ax.legend(fontsize=10)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, dpi=150, bbox_inches='tight', format='png')
        buf.seek(0)
        charts['owner_comparison'] = buf
        plt.close()

    # 4. 高价值机会Top20
    if not analyzer.high_value.empty:
        top = analyzer.high_value.head(20).copy()
        fig, ax = plt.subplots(figsize=(12, 7))
        names = [str(n)[:30] for n in top['业务机会名'].tolist()]
        nom = top['_amt'].values / 10000
        wgt = top['_weighted_amt'].values / 10000
        sc = top['_total_score'].values
        y = range(len(names))

        ax.barh([i - 0.2 for i in y], nom, 0.4, label='名义金额', color='#3498db', alpha=0.8)
        ax.barh([i + 0.2 for i in y], wgt, 0.4, label='加权金额', color='#27ae60', alpha=0.8)
        for i, s in enumerate(sc):
            ax.text(max(nom[i], wgt[i]) + max(nom.max(), wgt.max()) * 0.02, i,
                    f'{s:.0f}分', va='center', fontsize=10, color='#e74c3c', fontweight='bold')

        ax.set_yticks(y)
        ax.set_yticklabels(names, fontsize=10)
        ax.set_xlabel('金额(万元)', fontsize=12)
        ax.set_title('高价值机会Top20(含综合评分)', fontsize=14, fontweight='bold')
        ax.invert_yaxis()
        ax.legend(fontsize=10)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, dpi=150, bbox_inches='tight', format='png')
        buf.seek(0)
        charts['high_value_top20'] = buf
        plt.close()

    # 5. 有效业务机会甄别过程漏斗
    fig, ax = plt.subplots(figsize=(10, 5))
    fn = list(analyzer.有效业务机会_filters.keys())
    fc_vals = [len(d) for d in analyzer.有效业务机会_filters.values()]

    ax.fill_between(range(len(fn)), fc_vals, alpha=0.3, color='#3498db')
    ax.plot(range(len(fn)), fc_vals, 'o-', color='#2c3e50', linewidth=2, markersize=8)
    for i, (n, c) in enumerate(zip(fn, fc_vals)):
        ax.annotate(str(c), (i, c), textcoords="offset points",
                    xytext=(0, 12), ha='center', fontsize=10, fontweight='bold')

    ax.set_xticks(range(len(fn)))
    ax.set_xticklabels([n.split('｜')[-1] if '｜' in n else n for n in fn],
                       rotation=25, ha='right', fontsize=10)
    ax.set_ylabel('机会数量', fontsize=12)
    ax.set_title('有效业务机会甄别过程漏斗', fontsize=14, fontweight='bold')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()

    buf = io.BytesIO()
    plt.savefig(buf, dpi=150, bbox_inches='tight', format='png')
    buf.seek(0)
    charts['有效业务机会_funnel'] = buf
    plt.close()

    return charts

# ================================================================
# DOCX报告生成（补充管理行动建议和附录，自动保存到指定目录）
# ================================================================
def set_cell_shading(cell, color):
    cell._tc.get_or_add_tcPr().append( parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>') )

def set_run_font(run, font_name='宋体', size=None, bold=None, color=None):
    run.font.name = font_name
    r = run._element
    rPr = r.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font_name)
    if size is not None: run.font.size = size
    if bold is not None: run.bold = bold
    if color is not None: run.font.color.rgb = color

def remove_empty_paragraph_before_picture(doc):
    """清除文档中紧跟在分页符或标题后面的空段落，彻底解决图表导致的空白页问题"""
    paras_to_remove = []
    for i, para in enumerate(doc.paragraphs):
        is_break_or_heading = False
        # 检查是否为分页符段落
        for run in para.runs:
            if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                is_break_or_heading = True
                break
        # 检查是否为标题 (通过样式判断)
        if not is_break_or_heading and para.style.name.startswith('Heading'):
            is_break_or_heading = True
        if is_break_or_heading and i + 1 < len(doc.paragraphs):
            next_para = doc.paragraphs[i + 1]
            # 如果下一个段落完全没有文字且没有任何Run，说明是插入图片留下的空壳
            if next_para.text == '' and not next_para.runs:
                paras_to_remove.append(next_para._element)
    for p_element in paras_to_remove:
        p_element.getparent().remove(p_element)

def add_styled_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, font_name='宋体', size=Pt(8), bold=True, color=RGBColor(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2C3E50')
    for ri, rd in enumerate(data):
        for ci, val in enumerate(rd):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            set_run_font(r, font_name='宋体', size=Pt(7.5))
            if ri % 2 == 0:
                set_cell_shading(cell, 'EBF5FB')
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table

def generate_report(analyzer, charts, output_dir=r'D:\部门业务机会分析报告'):
    # 创建保存目录
    os.makedirs(output_dir, exist_ok=True)

    # 生成文件名
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    output_path = os.path.join(output_dir, f'销售漏斗分析报告_{timestamp}.docx')

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2); section.bottom_margin = Cm(2)
    section.left_margin = Cm(2); section.right_margin = Cm(2)

    # 修改默认样式的字体为宋体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

    # 标题部分
    for _ in range(4): doc.add_paragraph('')
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('销售漏斗智能分析报告')
    set_run_font(r, font_name='宋体', size=Pt(28), bold=True, color=RGBColor(44, 62, 80))
    st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run('基于有效业务机会模型与漏斗健康度理论')
    set_run_font(r, font_name='宋体', size=Pt(14), color=RGBColor(127, 140, 141))
    doc.add_paragraph('')
    dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = dp.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
    set_run_font(r, font_name='宋体', size=Pt(11), color=RGBColor(149, 165, 166))

    # 筛选信息提示
    diag = analyzer.read_diagnosis
    yf = diag.get('year_filter', {})
    if yf and yf.get('excluded', 0) > 0:
        doc.add_paragraph('')
        wp = doc.add_paragraph(); wp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = wp.add_run( f'📅 已筛选2026年签约项目：原始{yf["before"]}条，' 
                       f'排除{yf["excluded"]}条非2026年项目，保留{yf["after"]}条' )
        set_run_font(r, font_name='宋体', size=Pt(9), color=RGBColor(52, 152, 219))

    # 缺失列提示
    if analyzer.missing_cols:
        doc.add_paragraph('')
        wp2 = doc.add_paragraph(); wp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = wp2.add_run(f'⚠️ 未识别列：{", ".join(analyzer.missing_cols)}')
        set_run_font(r, font_name='宋体', size=Pt(9), color=RGBColor(231, 76, 60))

    # 数据概览
    doc.add_paragraph('')
    dp2 = doc.add_paragraph(); dp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = dp2.add_run( f'📊 当前分析：{diag.get("total_columns", 0)}列，' 
                   f'{analyzer.summary["total_count"]}条记录' )
    set_run_font(r, font_name='宋体', size=Pt(9), color=RGBColor(52, 152, 219))

    # 目录
    doc.add_page_break()
    h1 = doc.add_heading('目 录', level=1)
    for run in h1.runs: set_run_font(run, font_name='宋体')
    catalog_items = [
        '一、数据总览与核心指标',
        '二、漏斗健康度诊断',
        '三、无效业务机会识别与剔除',
        '四、有效业务机会甄别与高价值机会清单',
        '五、人员漏斗透视',
        '六、管理行动建议',
        '附录A：数据字典',
        '附录B：评分规则说明'
    ]
    for item in catalog_items:
        p = doc.add_paragraph()
        r = p.add_run(item)
        set_run_font(r, font_name='宋体', size=Pt(11))

    # 第一章：数据总览与核心指标
    doc.add_page_break()
    h1 = doc.add_heading('一、数据总览与核心指标', level=1)
    for run in h1.runs: set_run_font(run, font_name='宋体')

    h2 = doc.add_heading('1.1 核心数据指标', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    exc_yr = analyzer.summary.get('excluded_by_year', 0)
    s = analyzer.summary
    kpi = [
        ['原始记录数', f'{yf.get("before", 0)}条', 'Excel全部记录'],
        ['非2026年排除', f'{exc_yr}条', '预计签约日期非2026年'],
        ['本次分析数', f'{s["total_count"]}条', '2026年签约项目'],
        ['总预计签约金额', format_amount(s['total_amount']), '2026年项目金额合计'],
        ['无效机会数', f'{s["invalid_count"]}个', f'占比{s["invalid_count"] / max(s["total_count"], 1) * 100:.1f}%'],
        ['无效机会金额', format_amount(s['invalid_amount']), '被剔除金额'],
        ['有效机会数', f'{s["valid_count"]}个', f'占比{s["valid_count"] / max(s["total_count"], 1) * 100:.1f}%'],
        ['有效机会金额', format_amount(s['valid_amount']), '有效漏斗金额'],
        ['高价值机会数', f'{s["high_value_count"]}个', '有效业务机会甄别真机会'],
        ['高价值加权金额', format_amount(s['high_value_weighted']), '阶段概率折算收入'],
    ]
    add_styled_table(doc, ['指标项', '数值', '说明'], kpi, [4, 4, 8])

    h2 = doc.add_heading('1.2 漏斗阶段分布', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    stage_stats = s.get('stage_distribution', {})
    if stage_stats:
        ta = sum(v['amount'] for v in stage_stats.values())
        sd_rows = []
        for sn in STAGE_ORDER:
            if sn in stage_stats:
                v = stage_stats[sn]
                w = STAGE_WEIGHT_ORDERED.get(sn, 0)
                wa = v['amount'] * w
                pct = v['amount'] / max(ta, 1) * 100
                note = ''
                if pct > 30 and sn in ('客户立项', '技术认可'): note = '⚠️ 大肚腩'
                if sn == '签约准备' and pct < 5 and ta > 0: note = '⚠️ 出口过窄'
                sd_rows.append([sn, str(v['count']), format_amount(v['amount']),
                               f'{pct:.1f}%', format_amount(wa), note])
        add_styled_table(doc, ['阶段', '数量', '金额', '占比', '加权金额', '说明'],
                        sd_rows, [2.5, 2, 3, 2, 3, 4])

        # 插入阶段分布图表
        if 'stage_distribution' in charts:
            doc.add_picture(charts['stage_distribution'], width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 第二章：漏斗健康度诊断
    doc.add_page_break()
    h1 = doc.add_heading('二、漏斗健康度诊断', level=1)
    for run in h1.runs: set_run_font(run, font_name='宋体')

    h2 = doc.add_heading('2.1 漏斗形状诊断', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    si = []
    if stage_stats:
        ta = sum(v['amount'] for v in stage_stats.values())
        for sn in STAGE_ORDER:
            if sn in stage_stats:
                pct = stage_stats[sn]['amount'] / max(ta, 1) * 100
                if sn in ('客户立项', '技术认可') and pct > 25:
                    si.append(f'•「{sn}」占比{pct:.1f}%，项目堆积严重。')
                if sn == '签约准备' and pct < 5 and ta > 0:
                    si.append(f'•「签约准备」仅占{pct:.1f}%，短期订单不足。')
    if si:
        for issue in si:
            p = doc.add_paragraph()
            r = p.add_run(issue)
            set_run_font(r, font_name='宋体', size=Pt(10))
            if '⚠️' in issue: r.font.color.rgb = RGBColor(231, 76, 60)
    else:
        p = doc.add_paragraph('漏斗形状相对健康。')
        set_run_font(p.runs[0], font_name='宋体', size=Pt(10))

    h2 = doc.add_heading('2.2 漏斗流速诊断', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    stag_count = {}
    for rec in analyzer.invalid_records:
        if 'stagnant' in rec['reason_code']:
            stag_count[rec.get('stage', '未知')] = stag_count.get(rec.get('stage', '未知'), 0) + 1
    if stag_count:
        sd2 = []
        for stg, cnt in sorted(stag_count.items(), key=lambda x: -x[1]):
            risk = '🔴 极高' if cnt > 5 else ('🟡 高' if cnt > 2 else '🟢 中')
            sd2.append([stg, str(cnt), risk, '核实推进状态' if cnt <= 5 else '立即清理'])
        add_styled_table(doc, ['停滞阶段', '数量', '风险', '建议'], sd2, [3, 3, 3, 7])

    # 插入无效原因图表
    if 'invalid_reasons' in charts:
        doc.add_picture(charts['invalid_reasons'], width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 第三章：无效业务机会识别与剔除
    doc.add_page_break()
    h1 = doc.add_heading('三、无效业务机会识别与剔除', level=1)
    for run in h1.runs: set_run_font(run, font_name='宋体')

    h2 = doc.add_heading('3.1 无效机会判定标准', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    p = doc.add_paragraph('以下标准对每条记录逐一判定，命中任一即标记为"无效机会"。')
    set_run_font(p.runs[0], font_name='宋体')

    rd = []
    for i, item in enumerate(INVALID_REASON_MAP.items(), 1):
        val_str = str(item[1]) if not isinstance(item[1], str) else item[1]
        parts = val_str.split('：', 1)
        rd.append([f'R{i}', parts[0] if parts else val_str, parts[1] if len(parts) > 1 else ''])
    add_styled_table(doc, ['规则', '判定标准', '说明'], rd, [2, 5, 9])

    h2 = doc.add_heading('3.2 无效机会明细', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')
    p = doc.add_paragraph(f'共 {len(analyzer.invalid_records)} 条无效记录：')
    set_run_font(p.runs[0], font_name='宋体')

    rg = defaultdict(list)
    for rec in analyzer.invalid_records:
        rg[rec['reason_code']].append(rec)

    for rc, recs in rg.items():
        doc.add_heading(f'3.2.{list(rg.keys()).index(rc)+1} {INVALID_REASON_MAP.get(rc, rc)}', level=3)
        rd_rows = []
        for rec in recs[:50]:  # 最多显示50条
            rd_rows.append([
                rec['name'], rec['stage'], rec['owner'], format_amount(rec['amount']),
                format_date(rec['create_date']), format_date(rec['stage_mod_date'])
            ])
        if rd_rows:
            add_styled_table(doc, ['业务机会名', '阶段', '负责人', '金额', '创建日期', '阶段修改日期'],
                            rd_rows, [5, 2, 2, 2, 2, 2])

    # 第四章：有效业务机会甄别与高价值机会清单
    doc.add_page_break()
    h1 = doc.add_heading('四、有效业务机会甄别与高价值机会清单', level=1)
    for run in h1.runs: set_run_font(run, font_name='宋体')

    h2 = doc.add_heading('4.1 有效机会甄别流程', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')

    # 插入甄别漏斗图表
    if '有效业务机会_funnel' in charts:
        doc.add_picture(charts['有效业务机会_funnel'], width=Inches(7))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    h2 = doc.add_heading('4.2 高价值机会Top20', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')

    if not analyzer.high_value.empty:
        hv_rows = []
        for idx, row in analyzer.high_value.head(20).iterrows():
            hv_rows.append([
                idx+1, row['业务机会名'], row['阶段'], row['业务机会所有人'],
                format_amount(row['_amt']), format_amount(row['_weighted_amt']),
                f"{row['_total_score']:.0f}分"
            ])
        add_styled_table(doc, ['排名', '业务机会名', '阶段', '负责人', '金额', '加权金额', '综合评分'],
                        hv_rows, [1, 5, 2, 2, 2, 2, 2])

        # 插入高价值机会图表
        if 'high_value_top20' in charts:
            doc.add_picture(charts['high_value_top20'], width=Inches(7))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 第五章：人员漏斗透视
    doc.add_page_break()
    h1 = doc.add_heading('五、人员漏斗透视', level=1)
    for run in h1.runs: set_run_font(run, font_name='宋体')

    h2 = doc.add_heading('5.1 销售人员业绩分布', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')

    od = analyzer.summary.get('owner_distribution', {})
    if od:
        od_rows = []
        for name, stats in sorted(od.items(), key=lambda x: -x[1]['weighted'])[:20]:
            od_rows.append([
                name, stats['count'], format_amount(stats['amount']),
                format_amount(stats['weighted'])
            ])
        add_styled_table(doc, ['销售人员', '有效机会数', '名义金额', '加权金额'],
                        od_rows, [3, 2, 3, 3])

        # 插入人员对比图表
        if 'owner_comparison' in charts:
            doc.add_picture(charts['owner_comparison'], width=Inches(7))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 第六章：管理行动建议（补充）
    doc.add_page_break()
    h1 = doc.add_heading('六、管理行动建议', level=1)
    for run in h1.runs: set_run_font(run, font_name='宋体')

    h2 = doc.add_heading('6.1 紧急行动项（1周内）', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')

    urgent_actions = [
        '• 清理停滞项目：对签约准备阶段超过180天、技术/商务认可超过360天的项目进行逐一核实，确认是否继续跟进或直接关闭',
        '• 零金额/微型金额项目处理：对金额为0或小于5000元的项目进行筛选，无价值项目直接清理',
        '• 超期未更新项目跟进：对超过180天未更新的项目，要求销售人员在3个工作日内完成状态更新',
        '• 阶段概率校准：对签约准备阶段可能性低于90%的项目，重新评估阶段合理性'
    ]
    for action in urgent_actions:
        p = doc.add_paragraph()
        r = p.add_run(action)
        set_run_font(r, font_name='宋体', size=Pt(10), bold=True)
        r.font.color.rgb = RGBColor(220, 53, 69)

    h2 = doc.add_heading('6.2 短期改进项（1个月内）', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')

    short_actions = [
        '• 漏斗形状优化：针对"大肚腩"阶段（客户立项/技术认可占比过高），制定推进计划，加速项目流转',
        '• 销售人员赋能：对加权金额较低的销售人员进行一对一辅导，提升项目跟进质量',
        '• 数据质量提升：规范业务机会录入标准，确保必填字段（金额、预计签约日期、阶段修改日期）完整准确',
        '• 高价值项目聚焦：建立Top20高价值项目专项跟进机制，每周复盘推进进度'
    ]
    for action in short_actions:
        p = doc.add_paragraph()
        r = p.add_run(action)
        set_run_font(r, font_name='宋体', size=Pt(10))
        r.font.color.rgb = RGBColor(255, 193, 7)

    h2 = doc.add_heading('6.3 长期优化项（3个月内）', level=2)
    for run in h2.runs: set_run_font(run, font_name='宋体')

    long_actions = [
        '• 建立漏斗健康度监控体系：每周自动生成漏斗分析报告，监控各阶段转化率和流速',
        '• 完善阶段定义与考核机制：明确各销售阶段的判定标准和推进要求，将漏斗健康度纳入销售人员考核',
        '• 客户分层管理：基于高价值项目特征，建立客户分层体系，优化资源配置',
        '• 数字化工具升级：基于分析结果优化CRM系统，增加自动提醒和异常预警功能'
    ]
    for action in long_actions:
        p = doc.add_paragraph()
        r = p.add_run(action)
        set_run_font(r, font_name='宋体', size=Pt(10))
        r.font.color.rgb = RGBColor(40, 167, 69)

    # 附录A：数据字典
    doc.add_page_break()
    h1 = doc.add_heading('附录A：数据字典', level=1)
    for run in h1.runs: set_run_font(run, font_name='宋体')

    dict_rows = []
    for col in ALL_STANDARD_NAMES[:30]:  # 显示前30个核心字段
        desc = {
            '阶段': '业务机会当前所处的销售阶段',
            '预计签约金额': '项目预计签约的总金额',
            '业务机会所有人': '负责该项目的销售人员',
            '创建日期': '业务机会录入系统的日期',
            '预计签约日期': '项目预计完成签约的日期',
            '阶段修改日期': '最近一次阶段变更的日期',
            '可能性': '项目成功签约的概率（百分比）',
            '更新时间': '业务机会最近一次修改的时间'
        }.get(col, '无详细说明')
        dict_rows.append([col, desc])

    add_styled_table(doc, ['字段名', '字段说明'], dict_rows, [4, 10])

    # 附录B：评分规则说明
    doc.add_page_break()
    h1 = doc.add_heading('附录B：评分规则说明', level=1)
    for run in h1.runs: set_run_font(run, font_name='宋体')

    score_rules = [
        ['阶段得分（50分）', '根据销售阶段权重计算：销售线索2.5分、公司立项5分、客户立项15分、技术认可25分、商务认可37.5分、签约准备/已赢单50分'],
        ['金额得分（30分）', '5000-3万：5分；3-10万：10分；10-50万：20分；50万以上：30分'],
        ['类型得分（20分）', '直签用户：20分；其他类型：0分'],
        ['综合得分', '阶段得分 + 金额得分 + 类型得分（满分100分）']
    ]
    add_styled_table(doc, ['评分维度', '规则说明'], score_rules, [3, 11])

    # 保存文档
    doc.save(output_path)
    remove_empty_paragraph_before_picture(doc)  # 清理空段落
    doc.save(output_path)  # 重新保存

    return output_path

# ================================================================
# 主界面（移除下载按钮，自动保存报告）
# ================================================================
def main():
    root = tk.Tk()
    root.title('销售漏斗智能分析工具 v1.0')
    root.geometry('600x400')
    root.resizable(False, False)

    def select_file():
        file_path = filedialog.askopenfilename(
            title='选择Excel文件',
            filetypes=[('Excel文件', '*.xlsx *.xls *.csv'), ('所有文件', '*.*')]
        )
        if not file_path:
            return

        try:
            # 读取数据
            status_label.config(text='正在读取数据...')
            root.update()
            df, header = smart_read_excel(file_path)

            # 运行分析
            status_label.config(text='正在分析数据...')
            root.update()
            analyzer = FunnelAnalyzer(df)
            analyzer.run_analysis()

            # 生成图表（内存中）
            status_label.config(text='正在生成图表...')
            root.update()
            charts = generate_charts(analyzer)

            # 生成并保存报告
            status_label.config(text='正在生成报告...')
            root.update()
            report_path = generate_report(analyzer, charts)

            status_label.config(text=f'分析完成！报告已保存至：\n{report_path}')
            messagebox.showinfo('成功', f'分析完成！\n报告已保存至：\n{report_path}')

        except Exception as e:
            status_label.config(text=f'出错：{str(e)}')
            messagebox.showerror('错误', f'分析失败：{str(e)}')

    # 界面布局
    title_label = tk.Label(root, text='销售漏斗智能分析工具', font=('宋体', 18, 'bold'))
    title_label.pack(pady=30)

    select_btn = tk.Button(root, text='选择Excel文件开始分析', font=('宋体', 12),
                          width=30, height=2, command=select_file)
    select_btn.pack(pady=20)

    status_label = tk.Label(root, text='等待选择文件...', font=('宋体', 10),
                           fg='gray', justify=tk.CENTER)
    status_label.pack(pady=50)

    # 版权信息
    copyright_label = tk.Label(root, text='© 2026 销售分析部', font=('宋体', 9), fg='gray')
    copyright_label.pack(side=tk.BOTTOM, pady=10)

    root.mainloop()

if __name__ == '__main__':
    main()